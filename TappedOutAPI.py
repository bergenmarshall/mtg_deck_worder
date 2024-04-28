import requests
from time import sleep
import docx
import os
import shutil

name = "ezuri-claw-of-progress"

card_list = []

with open(f"{name}.txt", "r") as f:
    file_input = f.read().split("\n")

for card in file_input:
    num_of_that_card = card.split(" ")[0]
    card_name = card[2:].strip()
    for i in range(int(num_of_that_card)):
        card_list.append(card_name)

name_and_image = {}

print("Getting card images...")
for card in list(set(card_list)):
    card_url = card.replace(" ", "+")
    response = requests.get(f"https://api.scryfall.com/cards/named?fuzzy={card_url}")
    sleep(0.1)
    response = response.json()
    if len(response["multiverse_ids"]) > 1:
        print("More than one face found for", card)
        name_and_image[card] = [response["card_faces"][i]["image_uris"]["png"] for i in range(len(response["card_faces"]))]
    else: 
        name_and_image[card] = response["image_uris"]["png"]

if os.path.exists(name):
    shutil.rmtree(name)
os.mkdir(name)

print("Saving images...")
for card in name_and_image:
    file_name = card.replace(" ", "-")
    file_name = file_name.replace("/", "-")
    with open(f"{name}/{file_name}.png", "wb") as f:
        f.write(requests.get(name_and_image[card]).content)
    name_and_image[card] = f"{name}/{file_name}.png"

doc = docx.Document()
p = doc.add_paragraph()
r = p.add_run()

print("Adding images to document...")
for card in card_list:
    r.add_picture(name_and_image[card], width=docx.shared.Inches(2.5), height=docx.shared.Inches(3.5))
if os.path.exists(f"{name}.docx"):
    os.remove(f"{name}.docx")
doc.save(f"{name}.docx")

shutil.rmtree(name)